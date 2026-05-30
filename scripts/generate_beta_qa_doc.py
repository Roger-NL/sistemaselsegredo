from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_CELL_VERTICAL_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_BREAK
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


OUTPUT = "/Users/rogera/Documents/projetos/ES/configuracoes/QA_BETA_TESTERS_CHECKLIST.docx"

BLUE = RGBColor(0x2E, 0x74, 0xB5)
DARK_BLUE = RGBColor(0x1F, 0x4D, 0x78)
TEXT = RGBColor(0x22, 0x22, 0x22)
MUTED = RGBColor(0x66, 0x66, 0x66)
LIGHT_FILL = "E8EEF5"
LIGHTER_FILL = "F4F6F9"
BORDER = "C9D3E1"


def set_cell_shading(cell, fill):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def set_table_borders(table, color=BORDER, size="8"):
    tbl = table._tbl
    tbl_pr = tbl.tblPr
    borders = tbl_pr.first_child_found_in("w:tblBorders")
    if borders is None:
        borders = OxmlElement("w:tblBorders")
        tbl_pr.append(borders)
    for edge in ("top", "left", "bottom", "right", "insideH", "insideV"):
        el = borders.find(qn(f"w:{edge}"))
        if el is None:
            el = OxmlElement(f"w:{edge}")
            borders.append(el)
        el.set(qn("w:val"), "single")
        el.set(qn("w:sz"), size)
        el.set(qn("w:color"), color)


def set_cell_width(cell, inches):
    cell.width = Inches(inches)
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_w = tc_pr.find(qn("w:tcW"))
    if tc_w is None:
        tc_w = OxmlElement("w:tcW")
        tc_pr.append(tc_w)
    tc_w.set(qn("w:type"), "dxa")
    tc_w.set(qn("w:w"), str(int(inches * 1440)))


def configure_document(doc: Document):
    section = doc.sections[0]
    section.page_width = Inches(8.5)
    section.page_height = Inches(11)
    section.top_margin = Inches(1)
    section.bottom_margin = Inches(1)
    section.left_margin = Inches(1)
    section.right_margin = Inches(1)
    section.header_distance = Inches(0.49)
    section.footer_distance = Inches(0.49)

    styles = doc.styles
    normal = styles["Normal"]
    normal.font.name = "Calibri"
    normal.font.size = Pt(11)
    normal.font.color.rgb = TEXT
    normal.paragraph_format.space_after = Pt(6)
    normal.paragraph_format.line_spacing = 1.25

    for style_name, size, color in [
        ("Title", 22, TEXT),
        ("Heading 1", 16, BLUE),
        ("Heading 2", 13, BLUE),
        ("Heading 3", 12, DARK_BLUE),
    ]:
        style = styles[style_name]
        style.font.name = "Calibri"
        style.font.size = Pt(size)
        style.font.color.rgb = color
        style.paragraph_format.space_after = Pt(6)

    styles["Heading 1"].paragraph_format.space_before = Pt(18)
    styles["Heading 1"].paragraph_format.space_after = Pt(10)
    styles["Heading 2"].paragraph_format.space_before = Pt(14)
    styles["Heading 2"].paragraph_format.space_after = Pt(7)
    styles["Heading 3"].paragraph_format.space_before = Pt(10)
    styles["Heading 3"].paragraph_format.space_after = Pt(5)

    for sec in doc.sections:
        header = sec.header.paragraphs[0]
        header.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        run = header.add_run("QA Beta Launch | BasedSpeak")
        run.font.name = "Calibri"
        run.font.size = Pt(9)
        run.font.color.rgb = MUTED
        footer = sec.footer.paragraphs[0]
        footer.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        run = footer.add_run("Checklist operacional para testers")
        run.font.name = "Calibri"
        run.font.size = Pt(9)
        run.font.color.rgb = MUTED


def add_cover(doc: Document):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(140)
    r = p.add_run("QA COMPLETO DE BETA")
    r.font.name = "Calibri"
    r.font.size = Pt(30)
    r.font.bold = True
    r.font.color.rgb = TEXT

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_after = Pt(18)
    r = p.add_run("BasedSpeak | Fluxo real do produto antes da abertura para testers")
    r.font.name = "Calibri"
    r.font.size = Pt(15)
    r.font.color.rgb = DARK_BLUE

    meta = doc.add_table(rows=4, cols=2)
    meta.alignment = WD_TABLE_ALIGNMENT.CENTER
    meta.autofit = False
    set_table_borders(meta)
    rows = [
        ("Objetivo", "Validar o caminho real do aluno e do admin antes da abertura controlada para testers."),
        ("Escopo", "Cadastro, login, Pilar 1, convite, premium, Pilar 2, provas, agendamento, admin, regras e regressões."),
        ("Uso", "Marcar cada teste, registrar bug, classificar impacto e decidir go/no-go com base em evidência."),
        ("Versão", "Documento operacional multirodada para beta de credibilidade alta."),
    ]
    for row, (label, value) in zip(meta.rows, rows):
        set_cell_width(row.cells[0], 1.55)
        set_cell_width(row.cells[1], 4.95)
        row.cells[0].vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
        row.cells[1].vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
        set_cell_shading(row.cells[0], LIGHT_FILL)
        p0 = row.cells[0].paragraphs[0]
        p0.paragraph_format.space_after = Pt(0)
        rr = p0.add_run(label)
        rr.font.bold = True
        rr.font.size = Pt(10.5)
        rr.font.color.rgb = DARK_BLUE
        p1 = row.cells[1].paragraphs[0]
        p1.paragraph_format.space_after = Pt(0)
        rr = p1.add_run(value)
        rr.font.size = Pt(10.5)

    doc.add_paragraph()
    note = doc.add_paragraph()
    note.paragraph_format.space_before = Pt(18)
    note.paragraph_format.space_after = Pt(0)
    note.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = note.add_run("Uso recomendado: execute com duas contas de aluno, uma conta admin e anote cada falha no momento em que aparecer.")
    run.font.size = Pt(11)
    run.font.italic = True
    run.font.color.rgb = MUTED
    doc.add_page_break()


def add_section_title(doc, title, subtitle=None):
    p = doc.add_paragraph(title, style="Heading 1")
    p.runs[0].font.name = "Calibri"
    if subtitle:
        p2 = doc.add_paragraph(subtitle)
        p2.runs[0].font.size = Pt(10.5)
        p2.runs[0].font.color.rgb = MUTED


def add_bullets(doc, items):
    for item in items:
        p = doc.add_paragraph(style="Normal")
        p.paragraph_format.left_indent = Inches(0.38)
        p.paragraph_format.first_line_indent = Inches(-0.18)
        p.paragraph_format.space_after = Pt(4)
        r = p.add_run("• ")
        r.bold = True
        r.font.color.rgb = DARK_BLUE
        p.add_run(item)


def add_numbered(doc, items):
    for idx, item in enumerate(items, start=1):
        p = doc.add_paragraph(style="Normal")
        p.paragraph_format.left_indent = Inches(0.38)
        p.paragraph_format.first_line_indent = Inches(-0.18)
        p.paragraph_format.space_after = Pt(4)
        r = p.add_run(f"{idx}. ")
        r.bold = True
        r.font.color.rgb = DARK_BLUE
        p.add_run(item)


def add_matrix_table(doc, headers, rows, widths):
    table = doc.add_table(rows=1, cols=len(headers))
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = False
    set_table_borders(table)
    header_cells = table.rows[0].cells
    for cell, header, width in zip(header_cells, headers, widths):
        set_cell_width(cell, width)
        set_cell_shading(cell, LIGHT_FILL)
        p = cell.paragraphs[0]
        p.paragraph_format.space_after = Pt(0)
        r = p.add_run(header)
        r.font.bold = True
        r.font.size = Pt(10)
        r.font.color.rgb = DARK_BLUE
        cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER

    for row in rows:
        cells = table.add_row().cells
        for i, (value, width) in enumerate(zip(row, widths)):
            set_cell_width(cells[i], width)
            cells[i].vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
            p = cells[i].paragraphs[0]
            p.paragraph_format.space_after = Pt(0)
            r = p.add_run(value)
            r.font.size = Pt(10)

    return table


def add_test_block(doc, title, objective, checks):
    doc.add_paragraph(title, style="Heading 2")
    p = doc.add_paragraph()
    lead = p.add_run("Objetivo: ")
    lead.bold = True
    lead.font.color.rgb = DARK_BLUE
    p.add_run(objective)

    headers = ["Passo", "Ação de teste", "Resultado esperado", "Status"]
    widths = [0.65, 2.75, 2.55, 0.55]
    rows = []
    for idx, (action, expected) in enumerate(checks, start=1):
        rows.append((str(idx), action, expected, "☐"))
    add_matrix_table(doc, headers, rows, widths)
    doc.add_paragraph()


def add_bug_log_page(doc):
    add_section_title(doc, "Registro de Bugs e Triagem", "Use esta área durante a execução para não perder defeitos reais.")
    headers = ["ID", "Cenário", "Descrição do defeito", "Impacto", "Bloqueia?", "Responsável"]
    widths = [0.5, 1.3, 2.6, 0.8, 0.6, 0.7]
    rows = [("", "", "", "", "", "") for _ in range(10)]
    add_matrix_table(doc, headers, rows, widths)
    doc.add_paragraph("Critérios sugeridos:", style="Heading 2")
    add_bullets(doc, [
        "Bloqueador: impede cadastro, login, convite, premium, acesso correto ou operação admin essencial.",
        "Alto: produto ainda funciona, mas gera confiança baixa, estado inconsistente ou risco de suporte manual pesado.",
        "Médio: fluxo continua, porém com fricção, texto errado, erro visual importante ou mensagem ruim.",
        "Baixo: detalhe cosmético, microcópia, alinhamento visual ou comportamento secundário.",
    ])


def build_doc():
    doc = Document()
    configure_document(doc)
    add_cover(doc)

    add_section_title(doc, "Como usar este documento", "Rodada sugerida: primeiro smoke test, depois fluxo completo de aluno, depois admin e regressão.")
    add_numbered(doc, [
        "Separe as contas antes de começar: Conta A (aluno free novo), Conta B (aluno tester com convite), Conta C (admin).",
        "Execute os blocos na ordem do documento para reduzir falsos positivos causados por estado anterior.",
        "Sempre que um teste falhar, anote a rota, a conta usada, o passo exato, o comportamento esperado e o comportamento real.",
        "Se um item bloquear a continuação da rodada, marque como bloqueador e siga para outra conta ou cenário em vez de insistir no mesmo caminho.",
    ])
    add_bullets(doc, [
        "Ambiente recomendado: navegador limpo ou perfil anônimo para cada conta.",
        "Sempre recarregue a página depois de ações críticas: login, convite, aprovação de prova, promoção premium e agendamento.",
        "Anote também o que funcionou muito bem. Isso ajuda na decisão de go/no-go.",
    ])

    add_section_title(doc, "Contas e preparação", "Checklist de setup para começar a rodada sem ruído.")
    add_matrix_table(
        doc,
        ["Conta", "Função", "Pré-condição", "Checklist"],
        [
            ("Conta A", "Aluno free", "Nunca usada ou resetada", "☐"),
            ("Conta B", "Aluno tester", "Código de convite válido e novo", "☐"),
            ("Conta C", "Admin", "Acesso ao painel e aos códigos", "☐"),
            ("Ambiente", "Navegador", "Sem sessão antiga contaminando o teste", "☐"),
        ],
        [1.0, 1.45, 3.45, 0.6],
    )
    doc.add_page_break()

    add_section_title(doc, "Bloco 1 | Cadastro, login e persistência de sessão")
    add_test_block(doc, "1.1 Cadastro por e-mail", "Validar que a criação de conta termina em estado consistente de autenticação e perfil.", [
        ("Abrir a página de cadastro e criar uma conta nova com nome, e-mail e senha válidos.", "A conta é criada sem erro e a navegação pós-cadastro é coerente."),
        ("Sair da conta recém-criada e entrar novamente com as mesmas credenciais.", "O login funciona e o dashboard abre sem loop ou tela vazia."),
        ("Testar senha errada e e-mail inexistente.", "As mensagens de erro são claras e não travam a tela."),
    ])
    add_test_block(doc, "1.2 Callback e continuidade", "Garantir que uma rota protegida volta para o destino correto depois da autenticação.", [
        ("Tentar abrir uma rota protegida sem login, como /pilar/2 ou /agendamentos.", "A aplicação envia para login sem perder o destino original."),
        ("Fazer login a partir desse redirecionamento.", "Após autenticar, o usuário volta para a rota que tentou abrir."),
        ("Repetir o fluxo usando cadastro em vez de login.", "O callback continua válido até o fim do cadastro."),
    ])
    add_test_block(doc, "1.3 Perfil e persistência", "Confirmar que perfil e sessão sobrevivem a reload e nova navegação.", [
        ("Alterar nome e telefone no perfil.", "Os campos salvam e aparecem atualizados após recarga."),
        ("Fechar a aba, abrir novamente e voltar ao dashboard.", "A sessão continua íntegra ou pede login de forma coerente, sem estado quebrado."),
        ("Recarregar /dashboard, /perfil e /pagamento já logado.", "A experiência não entra em loading infinito nem inverte o estado logado/deslogado."),
    ])
    doc.add_page_break()

    add_section_title(doc, "Bloco 2 | Fluxo free até o fim do Pilar 1")
    add_test_block(doc, "2.1 Navegação e progresso do Pilar 1", "Validar que um aluno free novo consegue avançar pelo primeiro pilar sem inconsistência.", [
        ("Abrir o Pilar 1 com a Conta A e percorrer os módulos em ordem.", "Os módulos abrem corretamente e a progressão é salva."),
        ("Marcar avanço de módulos, sair e voltar ao Pilar 1.", "O progresso permanece salvo e coerente."),
        ("Concluir o último módulo e observar o próximo passo apresentado.", "A tela final do pilar mostra o estado certo, sem perguntas extras indevidas."),
    ])
    add_test_block(doc, "2.2 Fechamento do free", "Garantir que o aluno free termina o Pilar 1 e recebe o direcionamento correto.", [
        ("Após concluir o Pilar 1, voltar ao dashboard.", "O CTA premium/convite aparece no momento correto."),
        ("Tentar abrir o Pilar 2 ainda como free.", "O acesso é bloqueado do jeito esperado."),
        ("Abrir a página de pagamento a partir do CTA.", "A tela entende que o usuário é free e oferece o caminho correto."),
    ])
    add_bullets(doc, [
        "Anotar qualquer divergência entre o que o dashboard diz e o que a página do pilar diz.",
        "Anotar se o aluno parece ao mesmo tempo concluído e bloqueado em estados contraditórios.",
    ])
    doc.add_page_break()

    add_section_title(doc, "Bloco 3 | Convite, premium e acesso ao Pilar 2")
    add_test_block(doc, "3.1 Página de pagamento / convite", "Garantir que o tester com código entra pelo fluxo certo sem precisar pagar.", [
        ("Logar com a Conta B e abrir /pagamento.", "A página carrega sem erro e mostra claramente o campo de convite."),
        ("Testar um código inválido.", "A mensagem de erro é clara e não quebra a página."),
        ("Testar um código já usado.", "A resposta deixa claro que o código não está disponível."),
        ("Aplicar um código válido e novo.", "A conta vira premium e a transição visual faz sentido."),
    ])
    add_test_block(doc, "3.2 Persistência da virada premium", "Confirmar que premium não é só efeito visual momentâneo.", [
        ("Recarregar /pagamento depois do resgate do código.", "A página reconhece a conta já liberada."),
        ("Voltar ao dashboard e recarregar a página.", "O estado premium continua consistente."),
        ("Abrir /agendamentos.", "A rota premium abre corretamente."),
    ])
    add_test_block(doc, "3.3 Acesso ao Pilar 2", "Garantir que a promoção premium libera exatamente o que deve liberar.", [
        ("Abrir /pilar/2 com a Conta B após o resgate do convite.", "O Pilar 2 abre normalmente."),
        ("Tentar abrir pilares acima do que deveria estar liberado.", "Os pilares posteriores continuam obedecendo a progressão."),
        ("Abrir /especialidades.", "Continua bloqueado se essa ainda for a regra do produto."),
    ])
    doc.add_page_break()

    add_section_title(doc, "Bloco 4 | Provas, feedback e boletim")
    add_test_block(doc, "4.1 Envio de prova", "Validar o envio e o estado pendente do aluno.", [
        ("Enviar a prova do Pilar 1 em uma conta de aluno válida.", "A submissão é aceita e não cria erro de duplicidade indevida."),
        ("Voltar ao dashboard e ao boletim.", "O status pendente aparece corretamente."),
        ("Tentar reenviar a mesma prova enquanto ela está pendente.", "O sistema impede de forma clara e coerente."),
    ])
    add_test_block(doc, "4.2 Aprovação pelo admin", "Garantir que a avaliação muda o estado real do aluno.", [
        ("Entrar como admin e abrir a fila de aprovações.", "A prova pendente aparece na lista."),
        ("Aprovar a prova com feedback.", "A ação conclui sem erro e some/atualiza a fila corretamente."),
        ("Voltar à conta do aluno e recarregar dashboard, pilar e boletim.", "O novo estado aparece de forma consistente em todas as superfícies."),
    ])
    add_test_block(doc, "4.3 Reprovação pelo admin", "Conferir o comportamento difícil, não só o fluxo feliz.", [
        ("Reprovar uma prova com feedback.", "O estado reprovado aparece corretamente no aluno."),
        ("Voltar ao dashboard do aluno.", "A plataforma não tenta vender ou destravar algo fora de hora."),
        ("Abrir o boletim e o pilar correspondente.", "O feedback aparece e o próximo passo continua coerente."),
    ])
    doc.add_page_break()

    add_section_title(doc, "Bloco 5 | Agendamentos e progressão premium")
    add_test_block(doc, "5.1 Estado inicial", "Garantir que o aluno premium vê o estado correto de agendamento.", [
        ("Abrir /agendamentos em conta premium apta ao fluxo.", "A tela abre e mostra o estado real, não uma simulação quebrada."),
        ("Verificar dashboard e card de agendamento.", "As duas superfícies contam a mesma história."),
    ])
    add_test_block(doc, "5.2 Solicitação de horário", "Confirmar que o pedido do aluno entra sem inconsistência.", [
        ("Solicitar um horário, se a conta estiver no ponto de fazer isso.", "O pedido é salvo e o status muda."),
        ("Recarregar /agendamentos e /dashboard.", "O estado continua igual após reload."),
    ])
    add_test_block(doc, "5.3 Ação do admin", "Garantir que confirmação ou mudança de status refletem no aluno.", [
        ("Entrar no admin e aprovar/ajustar o agendamento.", "A ação administrativa conclui sem erro."),
        ("Voltar para a conta do aluno e recarregar.", "O novo estado aparece sem atraso estranho nem contradição."),
    ])
    doc.add_page_break()

    add_section_title(doc, "Bloco 6 | Painel admin e operação")
    add_test_block(doc, "6.1 Dashboard admin", "Conferir visão geral e métricas básicas.", [
        ("Abrir /admin/dashboard.", "Os cards carregam sem erro visível."),
        ("Confirmar contagem de usuários, códigos e sessões.", "Os dados parecem coerentes com o ambiente de teste."),
    ])
    add_test_block(doc, "6.2 Usuários e detalhe do aluno", "Validar leitura e ações operacionais do admin.", [
        ("Abrir /admin/usuarios e escolher um aluno.", "A lista abre e o detalhe carrega."),
        ("Ver provas, sessão ao vivo e dados de conta no detalhe.", "As informações aparecem sem permissão negada."),
        ("Alterar uma ação administrativa que você realmente usa.", "A alteração salva e reflete no aluno ou no próprio painel."),
    ])
    add_test_block(doc, "6.3 Códigos", "Garantir que o fluxo de testers é operável por você sem gambiarra.", [
        ("Abrir /admin/codigos.", "A lista abre."),
        ("Gerar um código novo.", "O código aparece e pode ser copiado/registrado."),
        ("Usar esse código em uma conta tester.", "O resgate funciona no fluxo real."),
    ])
    doc.add_page_break()

    add_section_title(doc, "Bloco 7 | Regressão, segurança aparente e leitura final")
    add_test_block(doc, "7.1 Rotas protegidas", "Checar o que não deveria abrir para conta errada.", [
        ("Conta free tentando abrir /pilar/2.", "Bloqueio correto."),
        ("Conta free tentando abrir /agendamentos.", "Bloqueio correto."),
        ("Conta comum tentando abrir rotas admin.", "Redirecionamento ou bloqueio correto."),
    ])
    add_test_block(doc, "7.2 Estado contraditório", "Caçar problemas de sessão e sincronização.", [
        ("Recarregar páginas críticas várias vezes seguidas.", "Nada entra em loop ou alterna entre estados contraditórios."),
        ("Trocar de conta e voltar rapidamente.", "A sessão limpa e carrega a conta correta."),
        ("Voltar ao dashboard depois de aprovações e convites.", "O estado final não depende de sorte ou cache velho."),
    ])
    add_test_block(doc, "7.3 Visual e microinterações", "Capturar os problemas que parecem pequenos, mas destroem confiança.", [
        ("Revisar home, login, cadastro, dashboard, pagamento, agendamentos e admin.", "Sem quebra visual séria, botão morto, modal preso ou texto cortado."),
        ("Prestar atenção em loading, mensagens de erro e empty states.", "Todos são compreensíveis e intencionais."),
    ])
    doc.add_page_break()

    add_bug_log_page(doc)
    doc.add_page_break()

    add_section_title(doc, "Go / No-Go da rodada", "Preencha ao final da sessão de testes.")
    add_matrix_table(
        doc,
        ["Item", "Pergunta de decisão", "Resposta / Observação"],
        [
            ("1", "Existe bloqueador em cadastro, login, convite, premium ou acesso ao Pilar 2?", ""),
            ("2", "Existe alguma ação admin essencial que falhou?", ""),
            ("3", "Existe contradição grave de estado entre dashboard, pilar, pagamento e agendamento?", ""),
            ("4", "Existe bug visual ou de confiança forte o bastante para prejudicar testers?", ""),
            ("5", "O produto está crível o suficiente para liberar a próxima leva de testers?", ""),
        ],
        [0.55, 3.75, 2.2],
    )
    doc.add_paragraph("Resumo final da rodada", style="Heading 2")
    for label in ["Bloqueadores", "Riscos altos", "Ajustes médios", "Ajustes cosméticos", "Decisão final"]:
        p = doc.add_paragraph()
        run = p.add_run(f"{label}: ")
        run.bold = True
        run.font.color.rgb = DARK_BLUE
        p.add_run("______________________________________________________________")

    doc.save(OUTPUT)


if __name__ == "__main__":
    build_doc()
